"""
CoreMatch — Document Service
Extracts text from PDF and DOCX files for AI pipeline processing.
Pure Python dependencies: PyPDF2 (PDF), python-docx (DOCX).
"""
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Supported MIME types
SUPPORTED_TYPES = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc',
}

MAX_CV_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


def extract_text(file_bytes: bytes, content_type: str) -> str:
    """Extract plain text from a PDF or DOCX file.

    Args:
        file_bytes: Raw file content.
        content_type: MIME type of the file.

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the document type is unsupported.
    """
    doc_type = SUPPORTED_TYPES.get(content_type)
    if doc_type == 'pdf':
        return _extract_pdf(file_bytes)
    elif doc_type in ('docx', 'doc'):
        return _extract_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported document type: {content_type}")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2."""
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    text = "\n".join(text_parts).strip()
    if not text:
        logger.warning("PDF text extraction produced empty result (possibly scanned/image PDF)")
    return text


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return "\n".join(text_parts).strip()


def validate_cv_file(file_bytes: bytes, content_type: str, filename: str) -> Optional[str]:
    """Validate a CV file for upload. Returns error message or None if valid."""
    if not file_bytes:
        return "File is empty"

    if len(file_bytes) > MAX_CV_SIZE_BYTES:
        return f"File too large. Maximum size is {MAX_CV_SIZE_BYTES // (1024 * 1024)}MB"

    if content_type not in SUPPORTED_TYPES:
        return f"Unsupported file type: {content_type}. Please upload a PDF or DOCX file"

    # Basic magic byte check
    if content_type == 'application/pdf' and not file_bytes[:5] == b'%PDF-':
        return "File does not appear to be a valid PDF"

    if content_type in ('application/vnd.openxmlformats-officedocument.wordprocessingml.document',) \
       and not file_bytes[:2] == b'PK':
        return "File does not appear to be a valid DOCX"

    return None
